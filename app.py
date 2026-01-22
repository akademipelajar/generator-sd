import streamlit as st
import json
import os
import time
import matplotlib
# Backend Agg wajib untuk stabilitas server
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from io import BytesIO
from urllib.parse import quote
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openai import OpenAI
from supabase import create_client

# ==========================================
# --- 1. INITIAL SETUP (WAJIB PALING ATAS) ---
# ==========================================
st.set_page_config(page_title="Generator Soal SD", page_icon="📚", layout="wide")

# Supabase Client Init
try:
    supabase = create_client(st.secrets["SUPABASE_URL"], st.secrets["SUPABASE_ANON_KEY"])
except Exception as e:
    st.error("Konfigurasi Supabase bermasalah. Cek Secrets.")
    st.stop()

# Initialize Session State
if "user" not in st.session_state:
    st.session_state.user = None
if 'hasil_soal' not in st.session_state: 
    st.session_state.hasil_soal = None
if 'reset_counter' not in st.session_state: 
    st.session_state.reset_counter = 0

# ==========================================
# --- 2. AUTHENTICATION GATE (HARD-LOCKED) ---
# ==========================================

def show_login_page():
    # Logika Recovery Mode
    q_params = st.query_params
    if "type" in q_params and q_params["type"] == "recovery":
        st.markdown("<br><br>", unsafe_allow_html=True)
        c1, c2, c3 = st.columns([1, 2, 1])
        with c2:
            st.title("🔄 Setel Password Baru")
            new_p = st.text_input("Password Baru", type="password", key="auth_recovery_pass")
            if st.button("Simpan Password Baru", key="auth_recovery_btn"):
                try:
                    supabase.auth.update_user({"password": new_p})
                    st.success("✅ Berhasil! Silakan login kembali.")
                    time.sleep(2)
                    st.query_params.clear()
                    st.rerun()
                except Exception as e: st.error(f"Gagal: {str(e)}")
        st.stop()

    st.markdown("<br><br>", unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 2, 1])
    with c2:
        st.title("🔐 Akses Akademi Pelajar")
        tab_l, tab_r, tab_f = st.tabs(["Login", "Daftar Akun", "Lupa Password"])
        
        with tab_l:
            l_email = st.text_input("Email", key="auth_l_email")
            l_pass = st.text_input("Password", type="password", key="auth_l_pass")
            if st.button("Masuk", key="auth_l_btn"):
                try:
                    res = supabase.auth.sign_in_with_password({"email": l_email, "password": l_pass})
                    st.session_state.user = res.user
                    st.rerun()
                except Exception as ex: st.error(f"Gagal login: {str(ex)}")
        
        with tab_r:
            r_email = st.text_input("Email Baru", key="auth_r_email")
            r_pass = st.text_input("Password Baru", type="password", key="auth_r_pass")
            if st.button("Daftar Sekarang", key="auth_r_btn"):
                try:
                    supabase.auth.sign_up({"email": r_email, "password": r_pass})
                    st.success("✅ Terdaftar! Silakan cek email atau login.")
                except Exception as ex: st.error(f"Gagal daftar: {str(ex)}")
        
        with tab_f:
            f_email = st.text_input("Email Terdaftar", key="auth_f_email")
            if st.button("Kirim Link Reset", key="auth_f_btn"):
                try:
                    supabase.auth.reset_password_for_email(f_email, {"redirect_to": "https://generator-sd.streamlit.app"})
                    st.success("📩 Link telah dikirim ke email Anda!")
                except Exception as ex: st.error(f"Gagal: {str(ex)}")
    st.stop()

# Pengecekan Login (Guard)
if st.session_state.user is None:
    try:
        session_res = supabase.auth.get_session()
        if session_res and session_res.session:
            st.session_state.user = session_res.user
        else:
            show_login_page()
    except:
        show_login_page()

# ==========================================
# --- 3. UI UTAMA (DESIGN LOCKED TOTAL) ---
# ==========================================

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=League+Spartan:wght@700&family=Poppins:ital,wght@1,700&display=swap');
    .header-title { font-family: 'League Spartan', sans-serif; font-size: 32px; font-weight: bold; line-height: 1.2; color: #1E1E1E; }
    .header-sub { font-family: 'Poppins', sans-serif; font-size: 18px; font-weight: bold; font-style: italic; color: #444; margin-bottom: 5px; }
    .warning-text { font-size: 13px; color: #d9534f; font-weight: bold; margin-bottom: 20px; }
    [data-testid="stSidebar"] { background: linear-gradient(180deg, #e6f3ff 0%, #ffffff 100%); border-right: 1px solid #d1e3f3; }
    .stRadio [data-testid="stWidgetLabel"] p, .stCheckbox p { font-weight: bold; font-size: 16px; color: #1E1E1E; }
    .metadata-text { font-size: 12px; font-style: italic; font-family: 'Poppins', sans-serif; font-weight: bold; color: #555; margin-top: 10px; margin-bottom: 15px;}
    div.stButton > button { width: 100%; }
    .table-header { background-color: #d1e3f3; padding: 10px; font-weight: bold; border: 1px solid #dee2e6; text-align: center; }
    .table-cell { padding: 10px; border: 1px solid #dee2e6; text-align: left; background-color: white; }
</style>
""", unsafe_allow_html=True)

# --- 4. DATABASE MATERI LENGKAP (MASTER RECOVERY - DIKUNCI) ---
DATABASE_MATERI = {
    "1 SD": {
        "Matematika": ["Bilangan sampai 10", "Penjumlahan & Pengurangan", "Bentuk Bangun Datar", "Mengukur Panjang Benda", "Mengenal Waktu"],
        "IPA": ["Anggota Tubuh", "Pancaindra", "Siang dan Malam", "Benda Hidup & Mati"],
        "Bahasa Indonesia": ["Mengenal Huruf", "Suku Kata", "Perkenalan Diri", "Benda di Sekitarku"],
        "Bahasa Inggris": ["Numbers 1-10", "Colors", "My Body", "Greetings"]
    },
    "2 SD": {
        "Matematika": ["Perkalian Dasar", "Pembagian Dasar", "Bangun Datar & Ruang", "Pengukuran Berat (kg, ons)", "Nilai Uang"],
        "IPA": ["Wujud Benda (Padat, Cair, Gas)", "Panas dan Cahaya", "Tempat Hidup Hewan"],
        "Bahasa Indonesia": ["Ungkapan Santun", "Tanda Baca", "Puisi Anak", "Menjaga Kesehatan"],
        "Bahasa Inggris": ["Parts of House", "Daily Activities", "Numbers 11-20", "Animals"]
    },
    "3 SD": {
        "Matematika": ["Pecahan Sederhana", "Simetri & Sudut", "Luas & Keliling Persegi", "Garis Bilangan", "Diagram Gambar"],
        "IPA": ["Ciri-ciri Makhluk Hidup", "Perubahan Wujud Benda", "Cuaca dan Musim", "Energi Alternatif"],
        "Bahasa Indonesia": ["Dongeng Hewan (Fabel)", "Lambang Pramuka", "Denah dan Arah", "Kalimat Saran"],
        "Bahasa Inggris": ["Telling Time", "Days of Week", "Hobby", "Professions"]
    },
    "4 SD": {
        "Matematika": ["Pecahan Senilai", "KPK dan FPB", "Segi Banyak", "Pembulatan Bilangan", "Diagram Batang"],
        "IPA": ["Gaya dan Gerak", "Sumber Energi", "Bunyi dan Pendengaran", "Cahaya", "Pelestarian Alam"],
        "Bahasa Indonesia": ["Gagasan Pokok", "Wawancara", "Puisi", "Teks Non-Fiksi"],
        "Bahasa Inggris": ["My Living Room", "Kitchen Utensils", "Present Continuous", "Transportation"]
    },
    "5 SD": {
        "Matematika": ["Operasi Pecahan", "Kecepatan dan Debit", "Skala dan Denah", "Volume Kubus & Balok", "Jaring-jaring Bangun Ruang"],
        "IPA": ["Organ Pernapasan", "Organ Pencernaan", "Ekosistem", "Panas dan Perpindahannya", "Siklus Air"],
        "Bahasa Indonesia": ["Iklan", "Pantun", "Surat Undangan", "Peristiwa Kebangsaan"],
        "Bahasa Inggris": ["Health Problems", "Ordering Food", "Comparison (Degrees)", "Shape and Size"]
    },
    "6 SD": {
        "Matematika": ["Bilangan Bulat Negatif", "Lingkaran", "Bangun Ruang Campuran", "Modus Median Mean", "Operasi Hitung Campuran"],
        "IPA": ["Perkembangbiakan Makhluk Hidup", "Listrik", "Magnet", "Tata Surya", "Pubertas"],
        "Bahasa Indonesia": ["Teks Laporan", "Pidato", "Formulir", "Teks Eksplanasi"],
        "Bahasa Inggris": ["Simple Past Tense", "Direction", "Government", "Holiday Plan"]
    }
}

LABEL_BENTUK = {
    "PG Sederhana": "Pilihlah satu jawaban yang benar",
    "PG Kompleks": "Pilihlah lebih dari satu jawaban yang benar",
    "PG Kompleks Kategori": "Pilih Benar atau Salah dari tiap pernyataan ini",
    "Uraian": "Jawablah pertanyaan berikut dengan tepat"
}

# --- 5. HELPER & WORD ENGINE ---
def get_clean_options(item):
    opsi_raw = item.get('opsi') or []
    labels = ['A', 'B', 'C', 'D']
    clean = []
    for i, text in enumerate(opsi_raw):
        if i >= 4: break
        t = str(text).strip()
        if t and not t.startswith(tuple(labels)): t = f"{labels[i]}. {t}"
        clean.append(t if t else f"{labels[i]}. [Kosong]")
    while len(clean) < 4: clean.append(f"{labels[len(clean)]}. [N/A]")
    return clean

def set_table_header_bg(cell):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'D1E3F3')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_docx(data_soal, mapel, kelas):
    doc = Document()
    doc.add_heading(f'LATIHAN SOAL {mapel.upper()}', 0)
    doc.add_paragraph(f'Kelas: {kelas}')
    doc.add_heading('A. DAFTAR SOAL', level=1)
    for idx, item in enumerate(data_soal):
        bentuk = item.get('bentuk', '')
        keterangan = LABEL_BENTUK.get(bentuk, "")
        p = doc.add_paragraph(); p.add_run(f"Soal {idx+1} ({keterangan})").italic = True
        doc.add_paragraph(item.get('soal',''), style='Normal').bold = True
        if bentuk == "PG Sederhana":
            for op in get_clean_options(item): doc.add_paragraph(op)
        elif bentuk == "PG Kompleks":
            for op in get_clean_options(item): doc.add_paragraph(f"☐ {op}")
        elif bentuk == "PG Kompleks Kategori":
            table = doc.add_table(rows=1, cols=3); table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = 'Pernyataan', 'Benar', 'Salah'
            for cell in hdr: set_table_header_bg(cell)
            for i, kat in enumerate(item.get('kategori_pernyataan', [])):
                row = table.add_row().cells
                row[0].text = f"{['A','B','C','D'][i]}. {kat['pernyataan']}"
        elif bentuk == "Uraian":
            doc.add_paragraph("Jawaban: ...................................................................")
        doc.add_paragraph(f"Materi : {item.get('materi','')} | Level : {item.get('level','')}")
    doc.add_page_break(); doc.add_heading('B. KUNCI JAWABAN & PEMBAHASAN', level=1)
    for idx, item in enumerate(data_soal):
        doc.add_paragraph(f"Nomor {idx+1}:").bold = True
        doc.add_paragraph(f"KUNCI: {item.get('kunci_jawaban_teks', '')}")
        doc.add_paragraph("PEMBAHASAN:")
        for step in item.get('pembahasan_langkah', []): doc.add_paragraph(f"• {step}")
        for analysis in item.get('analisis_opsi', []): doc.add_paragraph(f"• {analysis}")
        doc.add_paragraph(item.get('kesimpulan_akhir', '')).bold = True
        doc.add_paragraph("-" * 20)
    bio = BytesIO(); doc.save(bio); bio.seek(0); return bio

# --- 6. SIDEBAR (LOCKED) ---
with st.sidebar:
    suffix = st.session_state.reset_counter
    if os.path.exists("logo.png"):
        c1, c2, c3 = st.columns([1, 2, 1]); c2.image("logo.png", width=100)
    st.write(f"👤 **{st.session_state.user.email}**")
    if st.button("🚪 Logout", key="sidebar_logout_btn"):
        supabase.auth.sign_out(); st.session_state.user = None; st.rerun()
    st.divider()
    st.markdown("### ⚙️ Konfigurasi")
    api_key = st.secrets["OPENAI_API_KEY"]
    kelas_sel = st.selectbox("Pilih Kelas", list(DATABASE_MATERI.keys()), key=f"sel_k_{suffix}")
    mapel_sel = st.selectbox("Mata Pelajaran", list(DATABASE_MATERI[kelas_sel].keys()), key=f"sel_m_{suffix}")
    jml_soal = st.slider("Jumlah Soal", 1, 10, 2, key=f"sel_j_{suffix}")
    req_details = []
    for i in range(jml_soal):
        with st.expander(f"Soal {i+1}", expanded=(i==0)):
            top = st.selectbox("Materi", DATABASE_MATERI[kelas_sel][mapel_sel], key=f"t_{i}_{suffix}")
            lvl = st.selectbox("Level", ["Mudah", "Sedang", "Sulit (HOTS)"], key=f"l_{i}_{suffix}")
            fmt = st.selectbox("Bentuk Soal", list(LABEL_BENTUK.keys()), key=f"f_{i}_{suffix}")
            req_details.append({"topik": top, "level": lvl, "bentuk": fmt})
    c1, c2 = st.columns(2)
    btn_gen = c1.button("🚀 Generate", type="primary")
    if c2.button("🔄 Reset"):
        st.session_state.hasil_soal = None; st.session_state.reset_counter += 1; st.rerun()

# --- 7. MAIN PAGE HEADER ---
st.markdown('<div class="header-title">Generator Soal SD</div>', unsafe_allow_html=True)
st.markdown('<div class="header-sub">Berdasarkan Kurikulum Merdeka</div>', unsafe_allow_html=True)
st.write("---")

# --- 8. PERSONA MASTER & AI LOGIC (LOCKED & CUMULATIVE) ---
if btn_gen:
    client = OpenAI(api_key=api_key)
    status_box = st.status("✅ Soal Dalam Proses Pembuatan...", expanded=True)
    summary = "\n".join([f"- Soal {i+1}: {r['topik']}, {r['level']}, {r['bentuk']}" for i, r in enumerate(req_details)])
    
    # PERSONA MASTER: HOTS, Kemdikbud, Analisis Opsi, Kesimpulan, Korelasi
    system_prompt = """Anda adalah Pakar Pengembang Kurikulum Merdeka Kemdikbud RI dan Penulis Bank Soal Profesional SD. 
    Wajib memberikan jawaban dalam format json murni.

    KARAKTERISTIK HOTS: Mengukur C4-C6, Berpikir Kritis, Kontekstual Dunia Nyata.
    ATURAN KETAT:
    1. PG Sederhana: Pembahasan langkah demi langkah (pembahasan_langkah).
    2. PG Kompleks: Minimal 2 benar. Kunci wajib menyebutkan huruf (A, C, dst). Pembahasan per opsi (analisis_opsi).
    3. PG Kompleks Kategori: 4 Pernyataan label A, B, C, D. Jelaskan alasan Benar/Salah per label (analisis_opsi).
    4. Korelasi: Pertanyaan dan Jawaban/Pernyataan HARUS saling berhubungan erat secara logis.
    5. Kesimpulan Akhir: Wajib field 'kesimpulan_akhir' -> 'Jadi, jawaban yang benar adalah...'."""

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": f"Buat json soal SD Kurikulum Merdeka:\n{summary}"}],
            response_format={"type": "json_object"}
        )
        # Ambil data dan simpan ke state
        st.session_state.hasil_soal = json.loads(response.choices[0].message.content).get("soal_list", [])
        status_box.update(label="✅ Berhasil!", state="complete", expanded=False)
        st.rerun()
    except Exception as e: st.error(f"Gagal: {e}")

# --- 9. TAMPILAN HASIL (LOCKED DISPLAY LOGIC) ---
if st.session_state.hasil_soal:
    st.download_button("📥 Download Word", create_docx(st.session_state.hasil_soal, mapel_sel, kelas_sel), f"Soal_Master_{mapel_sel}.docx")
    for idx, item in enumerate(st.session_state.hasil_soal):
        with st.container(border=True):
            bentuk = item.get('bentuk'); ket = LABEL_BENTUK.get(bentuk)
            st.markdown(f"#### Soal {idx+1} *({ket})*")
            st.markdown(f"**{item.get('soal','')}**")
            
            if bentuk == "PG Sederhana": st.radio("Jawaban:", get_clean_options(item), key=f"ans_{idx}_{suffix}", index=None)
            elif bentuk == "PG Kompleks":
                for o_idx, opt in enumerate(get_clean_options(item)): st.checkbox(opt, key=f"chk_{idx}_{o_idx}_{suffix}")
            elif bentuk == "PG Kompleks Kategori":
                h1, h2, h3 = st.columns([4, 1, 1])
                h1.markdown("<div class='table-header'>Pernyataan</div>", unsafe_allow_html=True)
                h2.markdown("<div class='table-header'>Benar</div>", unsafe_allow_html=True)
                h3.markdown("<div class='table-header'>Salah</div>", unsafe_allow_html=True)
                lbls = ['A', 'B', 'C', 'D']
                for k_idx, kat in enumerate(item.get('kategori_pernyataan', [])):
                    c1, c2, c3 = st.columns([4, 1, 1])
                    c1.markdown(f"<div class='table-cell'>{lbls[k_idx]}. {kat['pernyataan']}</div>", unsafe_allow_html=True)
                    with c2: st.checkbox(" ", key=f"b_{idx}_{k_idx}_{suffix}", label_visibility="collapsed")
                    with c3: st.checkbox(" ", key=f"s_{idx}_{k_idx}_{suffix}", label_visibility="collapsed")
            elif bentuk == "Uraian": st.text_area("Jawaban:", key=f"txt_{idx}_{suffix}")

            # METADATA LOCKED
            st.markdown(f"<div class='metadata-text'>Materi : {item.get('materi','')} | Level : {item.get('level','')}</div>", unsafe_allow_html=True)
            with st.expander("Lihat Kunci & Pembahasan Mendalam"):
                kunci = item.get('kunci_jawaban_teks', '')
                if not kunci and bentuk == "PG Kompleks Kategori":
                    kunci = ", ".join([f"{lbls[i]}: {k['kunci']}" for i, k in enumerate(item.get('kategori_pernyataan', []))])
                st.success(f"**Kunci:** {kunci}")
                for s in item.get('pembahasan_langkah', []): st.write(f"✅ {s}")
                for a in item.get('analisis_opsi', []): st.write(f"• {a}")
                st.info(f"**Kesimpulan:** {item.get('kesimpulan_akhir','')}")

# --- 10. FOOTER (DIKUNCI TOTAL) ---
st.write("---")
st.markdown("<div style='text-align: center; font-size: 12px;'><b><p>Aplikasi Generator Soal ini Milik Bimbingan Belajar Digital \"Akademi Pelajar\"</p><p>Dilarang menyebarluaskan tanpa persetujuan tertulis dari Akademi Pelajar</p><p>Semua hak cipta dilindungi undang-undang</p></b></div>", unsafe_allow_html=True)
